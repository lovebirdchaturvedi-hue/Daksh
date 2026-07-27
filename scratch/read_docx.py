import docx
import sys

def read_docx(file_path):
    doc = docx.Document(file_path)
    output = []
    
    output.append("--- PARAGRAPHS ---")
    for i, p in enumerate(doc.paragraphs):
        output.append(f"P[{i}]: {p.text}")
        
    output.append("\n--- TABLES ---")
    for t_idx, table in enumerate(doc.tables):
        output.append(f"Table {t_idx}:")
        for r_idx, row in enumerate(table.rows):
            row_text = []
            for cell in row.cells:
                row_text.append(cell.text.replace("\n", " | "))
            output.append(f"  R[{r_idx}]: " + "\t".join(row_text))
            
    with open("docx_content.txt", "w", encoding="utf-8") as f:
        f.write("\n".join(output))
        
if __name__ == "__main__":
    read_docx("C:\\Users\\DELL\\Downloads\\Offer Letter-Ashish-1 (1) (1).docx")
